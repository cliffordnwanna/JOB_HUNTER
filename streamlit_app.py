# -*- coding: utf-8 -*-
"""
🚀 Remote Job Finder - Streamlit App v2.0
Find tailored remote jobs based on your CV
Enhanced with loading screen, better matching, and improved UX
"""

import streamlit as st
import io
import re
import time
import requests
import hashlib
from datetime import datetime, timedelta
from typing import Dict, List, Set, Optional
import pandas as pd
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings

warnings.filterwarnings("ignore")

# ============================================================
# PAGE CONFIG - Must be first Streamlit command
# ============================================================
st.set_page_config(
    page_title="🚀 Job Hunter",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="collapsed"  # Collapsed by default for mobile
)

# ============================================================
# LOADING SCREEN - Lightweight version
# ============================================================
# Quick initialization - no heavy model loading
if 'app_initialized' not in st.session_state:
    loading_placeholder = st.empty()
    loading_placeholder.markdown("""
    <style>
        .loading-container {
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            height: 60vh;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            border-radius: 20px;
            margin: 20px;
        }
        .loading-title {
            color: white;
            font-size: 2.5em;
            margin-bottom: 15px;
        }
        .loading-subtitle {
            color: #f0f0f0;
            font-size: 1.2em;
        }
    </style>
    <div class="loading-container">
        <div class="loading-title">🚀 Job Hunter</div>
        <div class="loading-subtitle">Ready to find your dream job!</div>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize app state
    st.session_state.app_initialized = True
    
    time.sleep(0.5)
    loading_placeholder.empty()

# ============================================================
# CUSTOM CSS - Mobile First Design
# ============================================================
st.markdown("""
<style>
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Mobile-friendly padding */
    .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
        padding-left: 1rem;
        padding-right: 1rem;
    }
    
    /* Prominent header */
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .main-header h1 {
        color: white;
        margin: 0;
        font-size: 2em;
        font-weight: 700;
    }
    .main-header p {
        color: #f0f0f0;
        margin: 8px 0 0 0;
        font-size: 1em;
    }
    
    /* Prominent filter card */
    .filter-card {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border: 2px solid #667eea;
        border-radius: 12px;
        padding: 1.2rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 2px 8px rgba(102, 126, 234, 0.15);
    }
    
    /* Job cards */
    .job-card {
        background: white;
        border: 1px solid #e0e0e0;
        border-left: 4px solid #667eea;
        padding: 1rem;
        border-radius: 8px;
        margin-bottom: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .job-title {font-size: 1.1em; font-weight: 600; color: #333; margin-bottom: 0.3rem;}
    .job-meta {font-size: 0.85em; color: #666;}
    .job-score {font-size: 1.4em; font-weight: bold;}
    .score-high {color: #4CAF50;}
    .score-medium {color: #FF9800;}
    .score-low {color: #999;}
    
    /* Mobile responsive */
    @media (max-width: 640px) {
        .stButton button {
            width: 100%;
            margin-bottom: 0.5rem;
        }
        .main-header h1 {
            font-size: 1.5em;
        }
        .main-header p {
            font-size: 0.9em;
        }
        .block-container {
            padding-left: 0.5rem;
            padding-right: 0.5rem;
        }
    }
    
    /* Success/info boxes */
    .stSuccess, .stInfo, .stWarning {
        border-radius: 8px;
        padding: 1rem;
    }
    
    /* Compact sidebar */
    section[data-testid="stSidebar"] {width: 280px !important;}
    section[data-testid="stSidebar"] .block-container {padding: 1rem;}
</style>
""", unsafe_allow_html=True)

# ============================================================
# CV PARSER
# ============================================================
class EnhancedCVParser:
    """Advanced CV parser with comprehensive skill extraction."""

    def __init__(self):
        self.cv_data = {}
        self.skill_database = {
            'social_media': [
                'instagram', 'facebook', 'twitter', 'tiktok', 'linkedin', 'youtube',
                'snapchat', 'pinterest', 'social media management', 'content creation',
                'content strategy', 'community management', 'social media marketing',
                'social media analytics', 'engagement', 'brand awareness', 'influencer marketing',
                'social media advertising', 'paid social', 'organic social', 'hashtag strategy',
                'social listening', 'crisis management', 'brand voice', 'copywriting',
                'content calendar', 'scheduling', 'hootsuite', 'buffer', 'sprout social',
                'later', 'meta business suite', 'facebook ads manager', 'instagram insights',
                'twitter analytics', 'tiktok analytics', 'canva', 'adobe creative suite',
                'photoshop', 'illustrator', 'video editing', 'premiere pro', 'final cut',
                'capcut', 'seo', 'sem', 'google analytics', 'social media reporting',
                'kpi tracking', 'roi analysis', 'a/b testing', 'audience insights',
                'social media strategy', 'brand management', 'reputation management'
            ],
            'data_science': [
                'python', 'r', 'sql', 'java', 'scala', 'javascript', 'c++',
                'machine learning', 'deep learning', 'neural networks', 'ai', 'artificial intelligence',
                'natural language processing', 'nlp', 'computer vision', 'cv',
                'data analysis', 'data science', 'statistics', 'probability',
                'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'sklearn',
                'pandas', 'numpy', 'scipy', 'matplotlib', 'seaborn', 'plotly',
                'jupyter', 'apache spark', 'hadoop', 'big data', 'etl',
                'aws', 'azure', 'gcp', 'google cloud', 'cloud computing',
                'docker', 'kubernetes', 'mlops', 'model deployment',
                'tableau', 'power bi', 'git', 'github', 'agile', 'scrum'
            ],
            'engineering': [
                'software engineering', 'web development', 'frontend', 'backend', 'full stack',
                'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi',
                'typescript', 'html', 'css', 'sass', 'tailwind', 'bootstrap',
                'mongodb', 'postgresql', 'mysql', 'redis', 'graphql', 'rest api',
                'ci/cd', 'jenkins', 'terraform', 'ansible', 'linux', 'devops',
                'microservices', 'system design', 'architecture', 'testing', 'unit testing'
            ],
            'design': [
                'ui design', 'ux design', 'user research', 'wireframing', 'prototyping',
                'figma', 'sketch', 'adobe xd', 'invision', 'zeplin',
                'graphic design', 'visual design', 'branding', 'typography', 'color theory',
                'motion graphics', 'animation', 'after effects', 'blender', '3d modeling',
                'product design', 'design systems', 'accessibility', 'responsive design'
            ],
            'business': [
                'sales', 'business development', 'account management', 'client relations',
                'negotiation', 'crm', 'salesforce', 'hubspot', 'pipedrive',
                'strategy', 'consulting', 'market research', 'competitive analysis',
                'financial analysis', 'budgeting', 'forecasting', 'excel', 'powerpoint',
                'operations', 'supply chain', 'logistics', 'procurement', 'vendor management'
            ],
            'hr_admin': [
                'recruiting', 'talent acquisition', 'onboarding', 'hr management',
                'payroll', 'benefits administration', 'employee relations', 'performance management',
                'administrative', 'executive assistant', 'office management', 'scheduling',
                'bookkeeping', 'quickbooks', 'invoicing', 'data entry', 'virtual assistant'
            ],
            'writing': [
                'content writing', 'technical writing', 'blog writing', 'article writing',
                'editing', 'proofreading', 'journalism', 'storytelling', 'creative writing',
                'grant writing', 'proposal writing', 'documentation', 'translation'
            ],
            'customer_service': [
                'customer support', 'customer service', 'help desk', 'technical support',
                'live chat', 'zendesk', 'intercom', 'freshdesk', 'ticketing',
                'conflict resolution', 'customer success', 'account management', 'retention'
            ],
            'soft_skills': [
                'communication', 'teamwork', 'leadership', 'project management',
                'problem solving', 'critical thinking', 'creativity', 'collaboration',
                'time management', 'organization', 'adaptability', 'remote work',
                'agile', 'scrum', 'cross-functional', 'stakeholder management',
                'presentation', 'writing', 'research', 'analytical', 'attention to detail'
            ]
        }

    def parse_pdf(self, file_content):
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                return text
        except Exception as e:
            raise ValueError(f"PDF parsing error: {str(e)}")

    def parse_docx(self, file_content):
        try:
            doc = Document(io.BytesIO(file_content))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            raise ValueError(f"DOCX parsing error: {str(e)}")

    def parse_txt(self, file_content):
        try:
            return file_content.decode('utf-8')
        except Exception as e:
            raise ValueError(f"TXT parsing error: {str(e)}")

    def extract_comprehensive_skills(self, text: str) -> Set[str]:
        if not isinstance(text, str):
            text = str(text)
        text_lower = text.lower()
        found_skills = set()

        all_skills = []
        for category in self.skill_database.values():
            all_skills.extend(category)

        for skill in all_skills:
            if skill in text_lower:
                found_skills.add(skill)

        prog_langs = ['python', 'java', 'javascript', 'r', 'sql', 'scala', 'c++', 'c#']
        for lang in prog_langs:
            if lang in text_lower:
                found_skills.add(lang)

        return found_skills

    def extract_info(self, text: str) -> Dict:
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)

        phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,9}'
        phones = re.findall(phone_pattern, text)

        found_skills = self.extract_comprehensive_skills(text)
        years_exp = self._extract_years_experience(text)

        return {
            'email': emails[0] if emails else 'Not found',
            'phone': phones[0] if phones else 'Not found',
            'skills': sorted(list(found_skills)),
            'years_experience': years_exp,
            'full_text': text,
            'skill_count': len(found_skills)
        }

    def _extract_years_experience(self, text: str) -> int:
        text_lower = text.lower()
        years = 0

        patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience[:\s]+(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s+(?:in|working|as)',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                years = max(years, max([int(m) for m in matches]))

        date_pattern = r'(20\d{2}|19\d{2})\s*[-–—]\s*(20\d{2}|19\d{2}|present|current)'
        date_ranges = re.findall(date_pattern, text_lower)

        if date_ranges:
            total_years = 0
            for start, end in date_ranges:
                start_year = int(start)
                end_year = 2026 if end in ['present', 'current'] else int(end)
                total_years += (end_year - start_year)
            years = max(years, total_years)

        return years

    def parse(self, filename: str, file_content: bytes) -> Dict:
        ext = filename.lower().split('.')[-1]

        if ext == 'pdf':
            text = self.parse_pdf(file_content)
        elif ext in ['docx', 'doc']:
            text = self.parse_docx(file_content)
        elif ext == 'txt':
            text = self.parse_txt(file_content)
        else:
            raise ValueError(f"Unsupported format: {ext}")

        self.cv_data = self.extract_info(text)
        return self.cv_data


# ============================================================
# JOB SCRAPER
# ============================================================
class JobScraper:
    """Job scraper using multiple sources."""

    def __init__(self):
        self.jobs = []

    def scrape_remoteok(self, limit=50):
        """Scrape RemoteOK - one of the largest remote job boards."""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            data = requests.get("https://remoteok.com/api", timeout=15, headers=headers).json()
            jobs = []
            for item in data[1:limit+1]:
                if isinstance(item, dict):
                    jobs.append({
                        'title': item.get('position', ''),
                        'company': item.get('company', ''),
                        'location': 'Remote',
                        'description': item.get('description', ''),
                        'tags': item.get('tags', []),
                        'url': f"https://remoteok.com/remote-jobs/{item.get('id', '')}",
                        'salary': item.get('salary_min', 'Not specified'),
                        'posted_date': item.get('date', 'N/A'),
                        'source': 'RemoteOK'
                    })
            return jobs
        except Exception as e:
            return []

    def scrape_remotive(self, limit=50):
        """Scrape Remotive - curated remote jobs in tech."""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get("https://remotive.com/api/remote-jobs", timeout=15, headers=headers)
            data = response.json()
            jobs = []
            for item in data.get('jobs', [])[:limit]:
                jobs.append({
                    'title': item.get('title', ''),
                    'company': item.get('company_name', ''),
                    'location': item.get('candidate_required_location', 'Remote') or 'Remote',
                    'description': item.get('description', ''),
                    'tags': [item.get('category', '')],
                    'url': item.get('url', ''),
                    'salary': item.get('salary', 'Not specified') or 'Not specified',
                    'posted_date': item.get('publication_date', 'N/A'),
                    'source': 'Remotive'
                })
            return jobs
        except Exception as e:
            print(f"Remotive error: {e}")
            return []

    def scrape_jobicy(self, limit=50):
        """Scrape Jobicy - remote jobs worldwide."""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get("https://jobicy.com/api/v2/remote-jobs?count=50", timeout=15, headers=headers)
            data = response.json()
            jobs = []
            for item in data.get('jobs', [])[:limit]:
                jobs.append({
                    'title': item.get('jobTitle', ''),
                    'company': item.get('companyName', ''),
                    'location': item.get('jobGeo', 'Remote') or 'Remote',
                    'description': item.get('jobDescription', ''),
                    'tags': [item.get('jobIndustry', '')],
                    'url': item.get('url', ''),
                    'salary': f"{item.get('annualSalaryMin', '')} - {item.get('annualSalaryMax', '')}".strip(' -') or 'Not specified',
                    'posted_date': item.get('pubDate', 'N/A'),
                    'source': 'Jobicy'
                })
            return jobs
        except Exception as e:
            print(f"Jobicy error: {e}")
            return []

    def scrape_arbeitnow(self, limit=50):
        """Scrape Arbeitnow - EU remote jobs."""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get("https://www.arbeitnow.com/api/job-board-api", timeout=15, headers=headers)
            data = response.json()
            jobs = []
            for item in data.get('data', []):
                if item.get('remote', False):  # Only remote jobs
                    jobs.append({
                        'title': item.get('title', ''),
                        'company': item.get('company_name', ''),
                        'location': item.get('location', 'Remote') or 'Remote',
                        'description': item.get('description', ''),
                        'tags': item.get('tags', []),
                        'url': item.get('url', ''),
                        'salary': 'Not specified',
                        'posted_date': item.get('created_at', 'N/A'),
                        'source': 'Arbeitnow'
                    })
                    if len(jobs) >= limit:
                        break
            return jobs
        except Exception as e:
            print(f"Arbeitnow error: {e}")
            return []

    def scrape_himalayas(self, limit=50):
        """Scrape Himalayas - remote-first company jobs."""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get("https://himalayas.app/jobs/api?limit=50", timeout=15, headers=headers)
            data = response.json()
            jobs = []
            for item in data.get('jobs', [])[:limit]:
                salary = 'Not specified'
                if item.get('minSalary'):
                    currency = item.get('salaryCurrency', '')
                    salary = f"{currency} {item.get('minSalary', '')}"
                jobs.append({
                    'title': item.get('title', ''),
                    'company': item.get('companyName', ''),
                    'location': 'Remote',
                    'description': item.get('description', ''),
                    'tags': item.get('categories', []) if isinstance(item.get('categories'), list) else [],
                    'url': f"https://himalayas.app/jobs/{item.get('slug', '')}",
                    'salary': salary,
                    'posted_date': item.get('pubDate', 'N/A'),
                    'source': 'Himalayas'
                })
            return jobs
        except Exception as e:
            print(f"Himalayas error: {e}")
            return []

    def scrape_weworkremotely(self, limit=50):
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            soup = BeautifulSoup(
                requests.get("https://weworkremotely.com/remote-jobs", timeout=15, headers=headers).text,
                "html.parser"
            )
            jobs = []
            for li in soup.select("li.feature")[:limit]:
                a = li.find("a", href=True)
                if not a:
                    continue
                jobs.append({
                    'title': li.select_one(".title").text.strip() if li.select_one(".title") else '',
                    'company': li.select_one(".company").text.strip() if li.select_one(".company") else '',
                    'location': 'Remote',
                    'description': '',
                    'tags': [],
                    'url': "https://weworkremotely.com" + a["href"],
                    'salary': 'Not specified',
                    'posted_date': 'Recent',
                    'source': 'WeWorkRemotely'
                })
            return jobs
        except Exception as e:
            st.warning(f"WeWorkRemotely unavailable: {e}")
            return []

    def scrape_jobspy(self, search_terms: List[str], country: str = "worldwide", results_per_site: int = 25):
        """Use python-jobspy to scrape LinkedIn, Indeed, Glassdoor (Google removed due to rate limiting)."""
        try:
            from jobspy import scrape_jobs
            
            search_query = " OR ".join(search_terms) if search_terms else "remote"
            
            # Note: google removed due to 429 rate limiting, zip_recruiter removed for EU/GDPR
            jobs_df = scrape_jobs(
                site_name=["indeed", "linkedin", "glassdoor"],
                search_term=search_query,
                location="remote",
                results_wanted=results_per_site,
                country_indeed=country if country != "worldwide" else "USA",
                hours_old=168,  # 7 days for more results
                is_remote=True
            )
            
            jobs = []
            for _, row in jobs_df.iterrows():
                jobs.append({
                    'title': str(row.get('title', '')),
                    'company': str(row.get('company', '')),
                    'location': str(row.get('location', 'Remote')),
                    'description': str(row.get('description', '')),
                    'tags': [],
                    'url': str(row.get('job_url', '')),
                    'salary': str(row.get('min_amount', 'Not specified')),
                    'posted_date': str(row.get('date_posted', 'N/A')),
                    'source': str(row.get('site', 'JobSpy'))
                })
            return jobs
        except ImportError:
            # JobSpy library not installed - silent fallback
            print("JobSpy not installed, using backup sources")
            return []
        except Exception as e:
            error_msg = str(e)
            # Handle common network errors gracefully
            if "NameResolutionError" in error_msg or "getaddrinfo failed" in error_msg:
                print(f"JobSpy: Network/DNS error - using backup sources")
            elif "Max retries exceeded" in error_msg:
                print(f"JobSpy: Rate limited or blocked - using backup sources")
            elif "timeout" in error_msg.lower():
                print(f"JobSpy: Timeout - using backup sources")
            else:
                print(f"JobSpy error: {error_msg[:100]}")
            return []

    def scrape_all(self, keywords: List[str] = None, progress_callback=None):
        """Scrape jobs from 9 different sources for maximum coverage."""
        all_jobs = []
        source_counts = {}

        # Source 1: JobSpy (LinkedIn, Indeed, Glassdoor)
        if progress_callback:
            progress_callback(0.05, "🔍 Searching LinkedIn, Indeed, Glassdoor...")
        jobspy_jobs = self.scrape_jobspy(keywords or ["remote"])
        source_counts['JobSpy'] = len(jobspy_jobs)
        all_jobs.extend(jobspy_jobs)

        # Source 2: RemoteOK
        if progress_callback:
            progress_callback(0.15, f"🔍 Searching RemoteOK... (JobSpy: {len(jobspy_jobs)})")
        remoteok_jobs = self.scrape_remoteok(limit=100)
        source_counts['RemoteOK'] = len(remoteok_jobs)
        all_jobs.extend(remoteok_jobs)

        # Source 3: Remotive
        if progress_callback:
            progress_callback(0.25, f"🔍 Searching Remotive... (Total: {len(all_jobs)})")
        remotive_jobs = self.scrape_remotive(limit=100)
        source_counts['Remotive'] = len(remotive_jobs)
        all_jobs.extend(remotive_jobs)

        # Source 4: Jobicy
        if progress_callback:
            progress_callback(0.35, f"🔍 Searching Jobicy... (Total: {len(all_jobs)})")
        jobicy_jobs = self.scrape_jobicy(limit=100)
        source_counts['Jobicy'] = len(jobicy_jobs)
        all_jobs.extend(jobicy_jobs)

        # Source 5: Arbeitnow (EU jobs)
        if progress_callback:
            progress_callback(0.45, f"🔍 Searching Arbeitnow... (Total: {len(all_jobs)})")
        arbeitnow_jobs = self.scrape_arbeitnow(limit=100)
        source_counts['Arbeitnow'] = len(arbeitnow_jobs)
        all_jobs.extend(arbeitnow_jobs)

        # Source 6: Himalayas
        if progress_callback:
            progress_callback(0.55, f"🔍 Searching Himalayas... (Total: {len(all_jobs)})")
        himalayas_jobs = self.scrape_himalayas(limit=100)
        source_counts['Himalayas'] = len(himalayas_jobs)
        all_jobs.extend(himalayas_jobs)

        # Source 7: WeWorkRemotely
        if progress_callback:
            progress_callback(0.65, f"🔍 Searching WeWorkRemotely... (Total: {len(all_jobs)})")
        wwr_jobs = self.scrape_weworkremotely(limit=100)
        source_counts['WeWorkRemotely'] = len(wwr_jobs)
        all_jobs.extend(wwr_jobs)

        # Log source counts to terminal
        print(f"\n📊 Job Source Results: {source_counts}")
        print(f"📊 Total before filtering: {len(all_jobs)}")

        if progress_callback:
            progress_callback(0.75, f"📊 Found {len(all_jobs)} jobs, filtering by relevance...")

        # Filter jobs by keyword relevance (title or description must contain at least one keyword)
        if keywords:
            keyword_filtered = []
            keywords_lower = [k.lower() for k in keywords]
            # Also include related terms for common searches
            expanded_keywords = set(keywords_lower)
            for kw in keywords_lower:
                if 'social media' in kw:
                    expanded_keywords.update(['social', 'media', 'content', 'marketing', 'community', 'digital'])
                if 'data' in kw:
                    expanded_keywords.update(['data', 'analyst', 'analytics', 'scientist', 'engineer'])
                if 'developer' in kw or 'engineer' in kw:
                    expanded_keywords.update(['developer', 'engineer', 'software', 'frontend', 'backend', 'fullstack'])
            
            for job in all_jobs:
                title_lower = job.get('title', '').lower()
                desc_lower = job.get('description', '').lower()[:500]  # First 500 chars
                combined = title_lower + ' ' + desc_lower
                
                # Check if any keyword matches
                if any(kw in combined for kw in expanded_keywords):
                    keyword_filtered.append(job)
            
            print(f"📊 After keyword filter: {len(keyword_filtered)} jobs")
            all_jobs = keyword_filtered

        if progress_callback:
            progress_callback(0.80, f"📊 Removing duplicates from {len(all_jobs)} jobs...")

        # Remove duplicates by URL and title+company combo
        seen_urls = set()
        seen_jobs = set()
        unique_jobs = []
        for job in all_jobs:
            url = job.get('url', '')
            job_key = f"{job.get('title', '').lower()}_{job.get('company', '').lower()}"
            
            if url and url not in seen_urls and job_key not in seen_jobs:
                seen_urls.add(url)
                seen_jobs.add(job_key)
                unique_jobs.append(job)

        print(f"📊 Final unique jobs: {len(unique_jobs)}")
        self.jobs = unique_jobs
        return unique_jobs


# ============================================================
# JOB MATCHER (LIGHTWEIGHT - TF-IDF based, no heavy dependencies)
# ============================================================
class JobMatcher:
    """Match jobs to CV using TF-IDF, skill matching, and role adjacency. Fast and reliable."""

    def __init__(self, cv_data: Dict):
        self.cv_data = cv_data
        self.cv_text = " ".join([
            " ".join(cv_data.get("skills", [])),
            cv_data.get("full_text", "")[:3000]  # Limit text size for performance
        ]).lower()

        self.vectorizer = TfidfVectorizer(stop_words="english", max_features=1000, ngram_range=(1, 2))

        # Universal role adjacency map - works for any profession
        self.adjacent_roles = {
            # Marketing & Social Media
            "social media": ["content creator", "community manager", "brand strategist",
                           "digital marketer", "marketing coordinator", "influencer"],
            "content": ["copywriter", "writer", "blogger", "editor", "content strategist"],
            "marketing": ["growth", "campaign manager", "brand coordinator", "seo"],
            # Tech & Engineering
            "software": ["developer", "engineer", "programmer", "full stack", "backend", "frontend"],
            "python": ["data scientist", "machine learning", "backend", "automation"],
            "javascript": ["frontend", "react", "node", "web developer", "full stack"],
            "data": ["analyst", "scientist", "bi", "reporting", "insights", "analytics"],
            # Design
            "design": ["ui", "ux", "graphic", "visual", "product designer", "creative"],
            "figma": ["ui designer", "ux designer", "product designer", "visual designer"],
            # Business & Sales
            "sales": ["business development", "account executive", "account manager", "bdm"],
            "account": ["customer success", "client relations", "account manager"],
            # Support & Admin
            "customer": ["support", "service", "help desk", "success", "chat"],
            "admin": ["assistant", "coordinator", "office manager", "executive assistant"],
            "virtual assistant": ["admin", "data entry", "scheduling", "bookkeeping"],
            # HR & Recruiting
            "recruiting": ["talent acquisition", "hr", "sourcer", "recruiter"],
            "hr": ["people operations", "talent", "employee relations"]
        }

        # Skill synonyms for better matching - universal
        self.skill_synonyms = {
            "python": ["programming", "coding", "development"],
            "javascript": ["js", "frontend", "web development"],
            "react": ["frontend", "ui development", "web app"],
            "sql": ["database", "data", "queries"],
            "excel": ["spreadsheet", "data analysis", "reporting"],
            "canva": ["graphic design", "visual design", "design tools"],
            "figma": ["ui design", "ux design", "prototyping"],
            "social media": ["smm", "social marketing", "digital marketing"],
            "content creation": ["content writing", "copywriting", "storytelling"],
            "customer service": ["customer support", "help desk", "client support"],
            "project management": ["pm", "coordination", "planning"],
            "analytics": ["data analysis", "reporting", "insights", "metrics"]
        }

    def expand_skills(self, skills: Set[str]) -> Set[str]:
        """Expand skills with synonyms for better matching."""
        expanded = set(skills)
        for skill in skills:
            if skill in self.skill_synonyms:
                expanded.update(self.skill_synonyms[skill])
        return expanded

    def match_job(self, job: Dict):
        title = job.get("Title", "").lower()
        desc = job.get("Description", "")[:1500].lower()  # Limit description size
        job_text = f"{title} {title} {desc}"  # Double-weight title

        # TF-IDF similarity (fast and reliable)
        tfidf_score = 0.0
        try:
            vectors = self.vectorizer.fit_transform([self.cv_text, job_text])
            tfidf_score = float(cosine_similarity(vectors[0:1], vectors[1:2])[0][0]) * 100
        except:
            tfidf_score = 0.0

        # Direct skill matching with synonyms
        cv_skills = set(s.lower() for s in self.cv_data.get("skills", []))
        expanded_skills = self.expand_skills(cv_skills)
        skill_matches = sum(1 for s in expanded_skills if s in job_text)
        skill_score = min((skill_matches / max(len(cv_skills), 1)) * 150, 100)  # Boost skill importance

        # Role adjacency bonus
        adjacency_bonus = 0
        for role_type, adjacent_list in self.adjacent_roles.items():
            if role_type in self.cv_text:
                for adj_role in adjacent_list:
                    if adj_role in job_text:
                        adjacency_bonus += 15
                        break

        # Title keyword bonus (direct match) - universal keywords
        title_bonus = 0
        important_keywords = [
            "manager", "coordinator", "specialist", "analyst", "developer", "engineer",
            "designer", "assistant", "associate", "consultant", "administrator",
            "marketing", "sales", "support", "content", "data", "software", "product"
        ]
        for kw in important_keywords:
            if kw in self.cv_text and kw in title:
                title_bonus += 8

        # Experience-based multiplier
        years = self.cv_data.get("years_experience", 0)
        entry_keywords = ["junior", "associate", "assistant", "coordinator", "entry", "intern", "specialist"]
        senior_keywords = ["senior", "lead", "head", "director", "manager", "principal"]
        
        is_entry = any(kw in title for kw in entry_keywords)
        is_senior = any(kw in title for kw in senior_keywords)
        
        if is_senior and years < 3:
            exp_multiplier = 0.7
        elif is_entry and 1 <= years <= 3:
            exp_multiplier = 1.2
        elif not is_senior and 1 <= years <= 4:
            exp_multiplier = 1.1
        else:
            exp_multiplier = 1.0

        # Final weighted score (TF-IDF + Skills + Bonuses)
        base_score = (
            (tfidf_score * 0.35) +
            (skill_score * 0.40) +
            adjacency_bonus +
            title_bonus
        )

        final_score = round(min(base_score * exp_multiplier, 100), 2)

        # Store results
        job["Match Score"] = final_score
        job["Match Explanation"] = f"Text:{tfidf_score:.0f}% | Skills:{skill_score:.0f}% | Fit:+{adjacency_bonus + title_bonus}"
        job["Skills Matched"] = skill_matches

        return final_score

    def score_jobs(self, jobs: List[Dict], progress_callback=None) -> List[Dict]:
        """Score all jobs - fast TF-IDF based matching."""
        total = len(jobs)
        for i, job in enumerate(jobs):
            self.match_job(job)
            if progress_callback and i % 10 == 0:
                progress_callback(0.6 + (0.3 * i / total), f"📊 Scoring job {i+1}/{total}...")
        
        return jobs


# ============================================================
# JOB FILTER
# ============================================================
class JobFilter:
    """Filter jobs based on recency, relevance, and preferences."""

    def __init__(self, preferences: Dict):
        self.preferences = preferences
        self.now = datetime.now()

    def parse_date(self, date_str: str) -> datetime:
        if not date_str or date_str in ['N/A', 'Recent']:
            return self.now
        try:
            if 'T' in str(date_str):
                dt = datetime.fromisoformat(str(date_str).replace('+00:00', '').replace('Z', ''))
                if dt.tzinfo:
                    dt = dt.replace(tzinfo=None)
                return dt
            return datetime.strptime(str(date_str)[:10], '%Y-%m-%d')
        except:
            return self.now

    def apply_all_filters(self, jobs_df: pd.DataFrame, max_days: int = 14, min_score: float = 50.0, exclude_keywords: List[str] = None, search_keywords: List[str] = None) -> pd.DataFrame:
        if len(jobs_df) == 0:
            return jobs_df

        # KEYWORD RELEVANCE BOOST - Boost jobs matching search terms (don't filter out others)
        if search_keywords:
            search_pattern = '|'.join([kw.strip().lower() for kw in search_keywords if kw.strip()])
            if search_pattern:
                title_lower = jobs_df['Title'].fillna('').str.lower()
                desc_lower = jobs_df['Description'].fillna('').str.lower()
                # Check if any search keyword appears in title or description
                title_match = title_lower.str.contains(search_pattern, na=False, regex=True)
                desc_match = desc_lower.str.contains(search_pattern, na=False, regex=True)
                jobs_df['keyword_match'] = title_match | desc_match
                # Boost score for keyword matches (+15 for title, +10 for description only)
                jobs_df.loc[title_match, 'Match Score'] = jobs_df.loc[title_match, 'Match Score'] + 15
                jobs_df.loc[desc_match & ~title_match, 'Match Score'] = jobs_df.loc[desc_match & ~title_match, 'Match Score'] + 10
                # Cap scores at 100%
                jobs_df['Match Score'] = jobs_df['Match Score'].clip(upper=100)
                # Don't filter - just boost. All jobs stay in results.

        # Exclude keywords filter - use word boundaries for accurate matching
        if exclude_keywords:
            for kw in exclude_keywords:
                kw = kw.strip().lower()
                if kw:
                    # Use word boundary regex for accurate matching
                    pattern = r'\b' + kw + r'\b'
                    title_lower = jobs_df['Title'].fillna('').str.lower()
                    jobs_df = jobs_df[~title_lower.str.contains(pattern, na=False, regex=True)].copy()

        # Recency filter
        jobs_df['parsed_date'] = jobs_df['Posted'].apply(self.parse_date)
        jobs_df['days_ago'] = (self.now - jobs_df['parsed_date']).dt.days
        filtered = jobs_df[jobs_df['days_ago'] <= max_days].copy()

        if len(filtered) < 20:
            filtered = jobs_df[jobs_df['days_ago'] <= 45].copy()

        # Score filter with progressive fallback to ensure results
        score_filtered = filtered[filtered['Match Score'] >= min_score].copy()
        if len(score_filtered) < 5:
            # Try lower threshold
            score_filtered = filtered[filtered['Match Score'] >= max(min_score - 20, 20)].copy()
        if len(score_filtered) < 5:
            # Try even lower
            score_filtered = filtered[filtered['Match Score'] >= 10].copy()
        if len(score_filtered) < 5:
            # Just take top jobs by score
            score_filtered = filtered.nlargest(min(20, len(filtered)), 'Match Score').copy()
        filtered = score_filtered

        # Location filter based on remote_type preference
        remote_type = self.preferences.get('remote_type', 'Remote (All)')
        location_text = filtered['Location'].fillna('').str.lower()
        
        if remote_type in ["Worldwide Remote (Anywhere)", "Worldwide (Anywhere)"]:
            # Detect US-specific locations (cities with state abbreviations like "New York, NY")
            us_state_pattern = r',\s*(al|ak|az|ar|ca|co|ct|de|fl|ga|hi|id|il|in|ia|ks|ky|la|me|md|ma|mi|mn|ms|mo|mt|ne|nv|nh|nj|nm|ny|nc|nd|oh|ok|or|pa|ri|sc|sd|tn|tx|ut|vt|va|wa|wv|wi|wy|dc)\b'
            is_us_city = location_text.str.contains(us_state_pattern, na=False, regex=True)
            
            # Also detect explicit US mentions
            us_country_pattern = r'\b(united states|usa|u\.s\.a|u\.s\b)'
            is_us_country = location_text.str.contains(us_country_pattern, na=False, regex=True)
            
            # Detect truly worldwide/remote jobs
            worldwide_keywords = ['anywhere', 'worldwide', 'global', 'international', 'emea', 'apac', 'latam', 'europe', 'africa', 'asia']
            worldwide_pattern = '|'.join(worldwide_keywords)
            is_worldwide = location_text.str.contains(worldwide_pattern, na=False, regex=True)
            
            # "Remote" without country restriction is good
            is_just_remote = location_text.str.strip() == 'remote'
            
            # Boost truly worldwide/remote jobs significantly
            filtered.loc[is_worldwide | is_just_remote, 'Match Score'] = filtered.loc[is_worldwide | is_just_remote, 'Match Score'] + 30
            
            # Penalize US-specific jobs heavily (cities or country) - they should appear lower
            is_us_specific = (is_us_city | is_us_country) & ~is_worldwide
            filtered.loc[is_us_specific, 'Match Score'] = filtered.loc[is_us_specific, 'Match Score'] - 40
                
        elif remote_type in ["USA Remote Only", "USA Only"]:
            # Keep only USA remote jobs
            usa_keywords = ['usa', 'us', 'united states', 'america']
            usa_pattern = '|'.join(usa_keywords)
            is_usa = location_text.str.contains(usa_pattern, na=False, regex=True)
            usa_jobs = filtered[is_usa].copy()
            if len(usa_jobs) >= 5:
                filtered = usa_jobs

        # Cap all scores at 100% to ensure no percentage exceeds 100
        filtered['Match Score'] = filtered['Match Score'].clip(upper=100)
        
        # Sort by score
        filtered = filtered.sort_values('Match Score', ascending=False).reset_index(drop=True)

        return filtered


# ============================================================
# APPLICATION HELPER (Enhanced with DOCX export)
# ============================================================
class ApplicationHelper:
    """Generate tailored application materials with DOCX export."""

    def extract_job_requirements(self, job_description: str) -> list:
        """Extract key requirements from job description."""
        if not job_description:
            return []
        desc_lower = job_description.lower()
        requirement_patterns = [
            'required', 'must have', 'should have', 'experience with',
            'knowledge of', 'proficiency in', 'expertise in', 'strong', 'excellent'
        ]
        requirements = []
        sentences = job_description.split('.')
        for sentence in sentences[:10]:
            sentence_lower = sentence.lower()
            if any(pattern in sentence_lower for pattern in requirement_patterns):
                clean = sentence.strip()
                if 20 < len(clean) < 150:
                    requirements.append(clean)
        return requirements[:5]

    def match_skills_to_job(self, cv_skills: list, job_description: str) -> list:
        """Find which CV skills match the job."""
        if not job_description:
            return cv_skills[:5]
        desc_lower = job_description.lower()
        matched_skills = [skill for skill in cv_skills if skill.lower() in desc_lower]
        return matched_skills[:8] if matched_skills else cv_skills[:5]

    def generate_cover_letter(self, cv_data: dict, job: dict, job_description: str = "") -> str:
        matched_skills = self.match_skills_to_job(cv_data.get('skills', []), job_description)
        skills_str = ', '.join(matched_skills[:5]) if matched_skills else ', '.join(cv_data.get('skills', [])[:5])
        
        # Detect role type for customization
        job_title_lower = job.get('Title', '').lower()
        is_data_role = any(term in job_title_lower for term in ['data', 'scientist', 'analyst', 'ai', 'machine learning', 'ml'])
        is_social_media = any(term in job_title_lower for term in ['social media', 'content', 'community', 'marketing'])
        
        if is_data_role:
            passion_line = "I am passionate about leveraging data-driven insights to solve complex problems and drive measurable business impact."
        elif is_social_media:
            passion_line = "I am passionate about creating engaging content that builds authentic connections and drives meaningful engagement."
        else:
            passion_line = "I am passionate about delivering high-quality results and contributing to team success."
        
        letter = f"""Dear Hiring Manager,

I am writing to express my strong interest in the {job.get('Title', 'Position')} position at {job.get('Company', 'your company')}. With {cv_data.get('years_experience', 1)}+ years of relevant experience and proven expertise in {skills_str}, I believe I would be an excellent fit for your team.

{passion_line} Throughout my career, I have developed strong capabilities in {', '.join(matched_skills[:3]) if matched_skills else ', '.join(cv_data.get('skills', [])[:3])}, which I understand are key requirements for this role.

What excites me most about this opportunity at {job.get('Company', 'your company')} is the chance to apply my skills in a remote environment where I can contribute meaningfully to your team's goals. I have a proven track record of successful remote collaboration and consistently delivering results in distributed team settings.

Key highlights of my qualifications include:
• {cv_data.get('years_experience', 1)}+ years of hands-on experience in relevant domains
• Strong proficiency in {matched_skills[0] if matched_skills else cv_data.get('skills', ['relevant tools'])[0]}
• Track record of delivering high-impact projects on time
• Excellent communication and collaboration skills in remote settings

I would welcome the opportunity to discuss how my background and skills align with your needs. I am confident that I can make valuable contributions to {job.get('Company', 'your team')}.

Thank you for considering my application.

Best regards,
[Your Name]
{cv_data.get('email', '[Your Email]')}
{cv_data.get('phone', '')}"""

        return letter

    def generate_resume_bullets(self, cv_data: dict, job: dict, job_description: str = "") -> list:
        """Generate tailored resume bullets."""
        matched_skills = self.match_skills_to_job(cv_data.get('skills', []), job_description)
        key_skills = matched_skills[:3] if matched_skills else cv_data.get('skills', [])[:3]
        
        bullets = [
            f"• Leveraged {key_skills[0] if key_skills else 'technical skills'}, {key_skills[1] if len(key_skills) > 1 else 'advanced tools'}, and {key_skills[2] if len(key_skills) > 2 else 'best practices'} to deliver high-impact solutions",
            f"• {cv_data.get('years_experience', 1)}+ years of progressive experience driving results in remote and distributed team environments",
            f"• Demonstrated expertise in {', '.join(matched_skills[:4]) if matched_skills else ', '.join(cv_data.get('skills', [])[:4])} with proven ability to adapt quickly",
            f"• Successfully managed multiple concurrent projects while maintaining high quality standards",
            f"• Strong communicator with experience presenting technical concepts to both technical and non-technical stakeholders"
        ]
        return bullets

    def generate_application_email(self, cv_data: dict, job: dict) -> str:
        email = f"""Subject: Application for {job.get('Title', 'Position')} Position - {cv_data.get('years_experience', 1)}+ Years Experience

Dear Hiring Team,

I am writing to apply for the {job.get('Title', 'Position')} position at {job.get('Company', 'your company')}. With {cv_data.get('years_experience', 1)}+ years of experience and a strong background in {', '.join(cv_data.get('skills', [])[:3])}, I am confident I would be a valuable addition to your team.

I am particularly drawn to this opportunity because of {job.get('Company', 'your company')}'s reputation and the exciting challenges this role presents. My experience in remote work environments has prepared me well for the collaborative and independent work this position requires.

I have attached my resume and cover letter for your review. I would welcome the opportunity to discuss how my skills and experience align with your needs.

Thank you for your time and consideration.

Best regards,
[Your Name]
{cv_data.get('email', '[Your Email]')}
{cv_data.get('phone', '')}
LinkedIn: [Your LinkedIn URL]"""

        return email

    def export_cover_letter_docx(self, cover_letter: str, job_title: str, company: str) -> bytes:
        """Export cover letter as DOCX and return bytes."""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(11)
        date_run.font.name = 'Calibri'
        
        doc.add_paragraph()
        
        # Add content
        for para_text in cover_letter.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.line_spacing = 1.15
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_all_materials_docx(self, cv_data: dict, job: dict, job_description: str = "") -> bytes:
        """Export complete application package as DOCX."""
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading(f"Application Materials for {job.get('Title', 'Position')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company info
        company_para = doc.add_paragraph(f"Company: {job.get('Company', 'Company')}")
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_para.runs[0].font.size = Pt(12)
        
        doc.add_page_break()
        
        # Section 1: Cover Letter
        doc.add_heading('Cover Letter', 1)
        cover_letter = self.generate_cover_letter(cv_data, job, job_description)
        for para_text in cover_letter.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.line_spacing = 1.15
        
        doc.add_page_break()
        
        # Section 2: Resume Bullets
        doc.add_heading('Tailored Resume Bullets', 1)
        doc.add_paragraph("Add these bullets to your resume when applying for this position:")
        doc.add_paragraph()
        bullets = self.generate_resume_bullets(cv_data, job, job_description)
        for bullet in bullets:
            para = doc.add_paragraph(bullet)
            para.paragraph_format.space_after = Pt(6)
        
        doc.add_page_break()
        
        # Section 3: Email Template
        doc.add_heading('Application Email Template', 1)
        email = self.generate_application_email(cv_data, job)
        for para_text in email.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.space_after = Pt(12)
        
        # Footer with job URL
        doc.add_paragraph()
        footer_para = doc.add_paragraph(f"Job URL: {job.get('URL', 'N/A')}")
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# ============================================================
# MAIN APP
# ============================================================
def main():
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🚀 Job Hunter</h1>
        <p>Find remote jobs that match your skills • Generate cover letters instantly</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Quick guide for new users
    st.info("💡 **How it works:** Upload your CV → Enter job titles → Get matched jobs → Apply with one click!")

    # Initialize session state
    if 'cv_data' not in st.session_state:
        st.session_state.cv_data = None
    if 'jobs_df' not in st.session_state:
        st.session_state.jobs_df = None
    if 'search_complete' not in st.session_state:
        st.session_state.search_complete = False
    if 'saved_jobs' not in st.session_state:
        st.session_state.saved_jobs = []
    if 'search_count' not in st.session_state:
        st.session_state.search_count = 0
    if 'show_materials' not in st.session_state:
        st.session_state.show_materials = False
    if 'selected_job_idx' not in st.session_state:
        st.session_state.selected_job_idx = 0
    # Filters in main area (permanently visible - no expander)
    st.subheader("⚙️ Search Filters")
    remote_type = st.selectbox(
        "🌍 Where can you work from?",
        ["Worldwide (Anywhere)", "All Remote", "USA Only"],
        index=0,
        help="Choose 'Worldwide' if you can work from any country"
    )
    min_match = st.slider("Minimum match score", 0, 100, 20, help="Lower = more results, Higher = better matches")
    max_days = st.slider("Job posted within", 1, 45, 14, help="How recent should the job postings be?")
    
    exclude_keywords = st.text_input(
        "Exclude job titles containing:",
        placeholder="senior, lead, director",
        help="Jobs with these words in the title will be hidden"
    )

    # Main content
    col1, col2 = st.columns([1, 1])

    with col1:
        st.header("📤 Step 1: Upload Your CV")
        uploaded_file = st.file_uploader(
            "Choose your CV file",
            type=['pdf', 'docx', 'doc', 'txt'],
            help="Supported formats: PDF, DOCX, TXT"
        )

        if uploaded_file is not None:
            if st.session_state.cv_data is None or st.session_state.get('uploaded_filename') != uploaded_file.name:
                with st.spinner("🔍 Parsing your CV..."):
                    parser = EnhancedCVParser()
                    file_content = uploaded_file.read()
                    st.session_state.cv_data = parser.parse(uploaded_file.name, file_content)
                    st.session_state.uploaded_filename = uploaded_file.name
                    st.session_state.search_complete = False

            cv_data = st.session_state.cv_data
            st.success(f"✅ CV Parsed! {cv_data['skill_count']} skills found.")

    with col2:
        st.header("🔍 Step 2: Search Jobs")
        
        job_titles = st.text_area(
            "What job are you looking for?",
            placeholder="e.g., social media manager, content creator, data analyst",
            help="Enter one or more job titles. Separate multiple titles with commas.",
            height=80
        )

        search_button = st.button("🚀 Find My Jobs", type="primary", use_container_width=True)

        if search_button:
            if st.session_state.cv_data is None:
                st.error("❌ Please upload your CV first!")
            elif not job_titles.strip():
                st.error("❌ Please enter at least one job title!")
            else:
                titles_list = [t.strip().lower() for t in job_titles.split(',') if t.strip()]
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                percent_text = st.empty()

                def update_progress(value, text):
                    progress_bar.progress(value)
                    percent = int(value * 100)
                    status_text.text(text)
                    percent_text.markdown(f"**{percent}%** complete")

                # Scrape jobs
                update_progress(0.05, "🔍 Starting job search...")
                scraper = JobScraper()
                jobs = scraper.scrape_all(keywords=titles_list, progress_callback=update_progress)

                if not jobs:
                    st.warning("⚠️ No jobs found. Try different keywords.")
                    st.session_state.search_complete = False
                else:
                    # Normalize jobs
                    jobs = [{
                        'Title': j.get('title', ''),
                        'Company': j.get('company', ''),
                        'Location': j.get('location', 'Remote'),
                        'Description': j.get('description', ''),
                        'URL': j.get('url', ''),
                        'Salary': j.get('salary', 'N/A'),
                        'Posted': j.get('posted_date', 'N/A'),
                        'Source': j.get('source', 'Unknown')
                    } for j in jobs]

                    update_progress(0.5, "🧠 Loading AI model...")
                    
                    # Score jobs
                    matcher = JobMatcher(st.session_state.cv_data)
                    scored_jobs = matcher.score_jobs(jobs, progress_callback=update_progress)
                    jobs_df = pd.DataFrame(scored_jobs)

                    update_progress(0.9, "🎯 Filtering results...")

                    # Filter jobs
                    job_filter = JobFilter({
                        'job_titles': titles_list,
                        'remote_type': remote_type
                    })
                    # Parse exclude keywords
                    exclude_list = [kw.strip() for kw in exclude_keywords.split(',') if kw.strip()] if exclude_keywords else []
                    # Pass search keywords for relevance filtering
                    filtered_df = job_filter.apply_all_filters(
                        jobs_df, 
                        max_days=max_days, 
                        min_score=min_match, 
                        exclude_keywords=exclude_list,
                        search_keywords=titles_list
                    )

                    st.session_state.jobs_df = filtered_df
                    st.session_state.search_complete = True
                    st.session_state.search_count += 1
                    st.session_state.jobs_shown = 15  # Reset pagination on new search

                    update_progress(1.0, "✅ Search complete!")
                    time.sleep(0.5)
                    progress_bar.empty()
                    status_text.empty()
                    percent_text.empty()
                    
                    # Auto-scroll to results
                    st.balloons()

    # Results section
    if st.session_state.search_complete and st.session_state.jobs_df is not None:
        st.divider()
        
        jobs_df = st.session_state.jobs_df
        
        if len(jobs_df) == 0:
            st.warning("⚠️ No matching jobs found. Try lowering the match threshold or broadening your search.")
            st.markdown("""
            <div class="tip-box">
                <strong>💡 Tips to find more jobs:</strong>
                <ul>
                    <li>Lower the minimum match % to 30-40%</li>
                    <li>Extend the date range to 30-45 days</li>
                    <li>Try broader job titles (e.g., "marketing" instead of "social media marketing")</li>
                    <li>Uncheck "Remote only" to see all positions</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        else:
            # Simple success message
            st.success(f"✅ Found **{len(jobs_df)} jobs** from {jobs_df['Source'].nunique()} sources")

            # Compact controls row
            col_sort, col_dl = st.columns([1, 1])
            with col_sort:
                sort_by = st.selectbox("Sort:", ["Best Match", "Most Recent"], key="sort_option", label_visibility="collapsed")
            with col_dl:
                csv = jobs_df.to_csv(index=False)
                st.download_button("📥 Download CSV", csv, f"jobs_{datetime.now().strftime('%Y%m%d')}.csv", "text/csv", use_container_width=True)
            
            # Apply sorting
            if sort_by == "Most Recent":
                jobs_df = jobs_df.sort_values('days_ago', ascending=True).reset_index(drop=True)
            else:
                jobs_df = jobs_df.sort_values('Match Score', ascending=False).reset_index(drop=True)

            # Pagination
            if 'jobs_shown' not in st.session_state:
                st.session_state.jobs_shown = 15
            
            jobs_to_show = min(st.session_state.jobs_shown, len(jobs_df))

            # Clean job list
            for idx, row in jobs_df.head(jobs_to_show).iterrows():
                score = row['Match Score']
                score_emoji = "🔥" if score >= 70 else "⭐" if score >= 50 else "📌"
                
                # Format posted date
                days_ago = row.get('days_ago', 0)
                if days_ago == 0:
                    posted_display = "Today"
                elif days_ago == 1:
                    posted_display = "Yesterday"
                elif days_ago <= 7:
                    posted_display = f"{int(days_ago)}d ago"
                else:
                    posted_display = "Recent"
                
                # Clean title (remove HTML entities)
                title = str(row['Title']).replace('&#8211;', '–').replace('&amp;', '&').replace('&#39;', "'")
                company = str(row['Company']) if row['Company'] and str(row['Company']).lower() not in ['nan', 'none', ''] else 'Company'
                loc = str(row['Location']) if row['Location'] else 'Remote'
                location = 'Remote' if loc.lower() in ['nan', 'none', ''] else loc[:30]
                
                with st.container():
                    st.markdown(f"**{score_emoji} {title}**")
                    st.caption(f"🏢 {company} • 📍 {location} • {posted_display}")
                    
                    # Action buttons row
                    btn1, btn2, btn3 = st.columns([2, 2, 1])
                    with btn1:
                        st.link_button("🔗 Apply Now", row['URL'], use_container_width=True, type="primary")
                    with btn2:
                        # Generate and download cover letter directly
                        helper = ApplicationHelper()
                        job_dict = {
                            'Title': row['Title'],
                            'Company': row['Company'],
                            'URL': row['URL'],
                            'Location': row.get('Location', 'Remote'),
                            'Description': row.get('Description', '')
                        }
                        cover_letter = helper.generate_cover_letter(st.session_state.cv_data, job_dict, job_dict.get('Description', ''))
                        docx_bytes = helper.export_cover_letter_docx(cover_letter, job_dict['Title'], job_dict['Company'])
                        st.download_button(
                            "📝 Cover Letter",
                            docx_bytes,
                            f"cover_letter_{job_dict['Company'].replace(' ', '_')}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key=f"dl_{idx}"
                        )
                    with btn3:
                        st.markdown(f"**{int(score)}%**")
                    st.divider()
            
            # Load More button and Generate Materials button
            if jobs_to_show < len(jobs_df):
                remaining = len(jobs_df) - jobs_to_show
                col1, col2 = st.columns(2)
                with col1:
                    if st.button(f"📄 Load More ({remaining})", use_container_width=True):
                        st.session_state.jobs_shown += 15
                        st.rerun()
                with col2:
                    if st.button("📋 Generate Application Materials", use_container_width=True, type="secondary"):
                        st.session_state.show_materials = True
                        st.rerun()
            else:
                # Full width button when no more jobs to load
                if st.button("📋 Generate Application Materials", use_container_width=True, type="secondary"):
                    st.session_state.show_materials = True
                    st.rerun()

    # Application Materials Generator Section
    if st.session_state.get('show_materials') and st.session_state.jobs_df is not None and len(st.session_state.jobs_df) > 0:
        st.divider()
        st.subheader("📝 Application Materials Generator")
        
        jobs_df = st.session_state.jobs_df
        
        col_select, col_close = st.columns([4, 1])
        
        # Use pre-selected job index if available
        selected_idx = st.session_state.get('selected_job_idx', 0)
        if selected_idx >= len(jobs_df):
            selected_idx = 0
        
        # Create job options with proper indexing
        job_options = [f"{i+1}. {row['Title'][:40]} @ {row['Company'][:20]}" for i, (idx, row) in enumerate(jobs_df.iterrows())]
        
        with col_select:
            selected_job = st.selectbox("Select a job to generate materials for:", job_options, index=selected_idx, key="job_selector")
        with col_close:
            if st.button("❌ Close", key="close_materials"):
                st.session_state.show_materials = False
                st.rerun()
        
        # Get the actual job index from the selected option
        job_idx = job_options.index(selected_job) if selected_job in job_options else selected_idx
        job_row = jobs_df.iloc[job_idx]
        job_dict = {
            'Title': job_row['Title'],
            'Company': job_row['Company'],
            'URL': job_row['URL'],
            'Location': job_row.get('Location', 'Remote'),
            'Description': job_row.get('Description', '')
        }
        
        st.info(f"**{job_dict['Title']}** at **{job_dict['Company']}** • {job_dict['Location']}")
        
        helper = ApplicationHelper()
        
        # Generate materials
        tab1, tab2, tab3 = st.tabs(["📧 Cover Letter", "📋 Resume Bullets", "✉️ Email Template"])
        
        with tab1:
            cover_letter = helper.generate_cover_letter(st.session_state.cv_data, job_dict, job_dict.get('Description', ''))
            st.text_area("Cover Letter", cover_letter, height=400, key="cover_letter_text")
            
            # Download as DOCX
            docx_bytes = helper.export_cover_letter_docx(cover_letter, job_dict['Title'], job_dict['Company'])
            st.download_button(
                "📥 Download Cover Letter (DOCX)",
                docx_bytes,
                f"cover_letter_{job_dict['Company'].replace(' ', '_')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with tab2:
            bullets = helper.generate_resume_bullets(st.session_state.cv_data, job_dict, job_dict.get('Description', ''))
            st.markdown("**Add these tailored bullets to your resume:**")
            for bullet in bullets:
                st.markdown(bullet)
        
        with tab3:
            email = helper.generate_application_email(st.session_state.cv_data, job_dict)
            st.text_area("Application Email", email, height=300, key="email_text")
        
        # Download complete package
        st.divider()
        all_materials_docx = helper.export_all_materials_docx(st.session_state.cv_data, job_dict, job_dict.get('Description', ''))
        st.download_button(
            "📦 Download Complete Application Package (DOCX)",
            all_materials_docx,
            f"application_package_{job_dict['Company'].replace(' ', '_')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )

    # Simple footer
    st.caption("📡 LinkedIn • Indeed • Glassdoor • RemoteOK • Remotive • Jobicy • Arbeitnow • Himalayas • WeWorkRemotely")
    
    # Hidden debug info (only shows if there's an issue)
    with st.expander("🔧 Debug Info", expanded=False):
        st.caption(f"Session ID: {id(st.session_state)}")
        st.caption(f"CV Loaded: {'Yes' if st.session_state.cv_data else 'No'}")
        st.caption(f"Jobs Found: {len(st.session_state.jobs_df) if st.session_state.jobs_df is not None else 0}")
        st.caption(f"Matching: TF-IDF + Skill-based (Fast Mode)")


if __name__ == "__main__":
    main()
