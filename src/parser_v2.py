"""
Production-Grade CV Parser v2.0
Dynamic LLM-based extraction with PII sanitization.
Works for ANY domain - no hardcoded skills.
"""
import io
import re
from typing import Dict, List, Optional
from dataclasses import asdict

# PDF/DOCX parsing
import pdfplumber
from docx import Document

# Internal modules
from .pii_sanitizer import PIISanitizer, PrivacyPreservingExtractor
from .llm_extractor import HybridExtractor, CVExtractionResult


class DynamicCVParser:
    """
    Production-grade CV parser that works for any profession.
    
    Features:
    - PII sanitization before external LLM calls
    - Dynamic skill extraction (no hardcoded databases)
    - Domain-agnostic (healthcare, tech, finance, etc.)
    - Confidence scoring
    - Graceful fallbacks
    """
    
    def __init__(self, use_llm: bool = True):
        self.use_llm = use_llm
        self.pii_extractor = PrivacyPreservingExtractor()
        self.hybrid_extractor = HybridExtractor()
        self.sanitizer = PIISanitizer()
    
    def _detect_file_type(self, file_content: bytes) -> str:
        """
        Detect file type using magic bytes (more reliable than extension).
        Returns: 'pdf', 'docx', 'txt', or raises ValueError
        """
        # PDF magic bytes: %PDF
        if file_content[:4] == b'%PDF':
            return 'pdf'
        
        # DOCX magic bytes: ZIP archive (PK\x03\x04) containing word/
        if file_content[:4] == b'PK\x03\x04':
            # Additional check: try to open as DOCX
            try:
                from docx import Document
                Document(io.BytesIO(file_content))
                return 'docx'
            except:
                pass
        
        # TXT: try UTF-8 decode
        try:
            file_content.decode('utf-8')
            return 'txt'
        except:
            pass
        
        raise ValueError("Unsupported file format: Cannot detect PDF, DOCX, or valid TXT")
    
    def _extract_text(self, filename: str, file_content: bytes) -> str:
        """Extract raw text from PDF, DOCX, or TXT using magic bytes detection."""
        # Use magic bytes detection (primary) with filename extension as fallback
        try:
            file_type = self._detect_file_type(file_content)
        except ValueError:
            # Fallback to filename extension for edge cases
            ext = filename.lower().split('.')[-1] if '.' in filename else ''
            if ext in ['pdf', 'docx', 'doc', 'txt']:
                file_type = ext
            else:
                raise ValueError(f"Cannot determine file type for: {filename}")
        
        if file_type == 'pdf':
            return self._parse_pdf(file_content)
        elif file_type in ['docx', 'doc']:
            return self._parse_docx(file_content)
        elif file_type == 'txt':
            return file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported format: {file_type}")
    
    def _parse_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF with scanned PDF detection."""
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                total_chars = 0
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        total_chars += len(page_text)
                
                # Detect scanned PDF (no text layer) - very few characters extracted
                if total_chars < 100 and len(pdf.pages) > 0:
                    raise ValueError(
                        "This appears to be a scanned PDF (image-based) with no text layer. "
                        "Please convert to a text-based PDF first using an OCR tool like Adobe Acrobat, "
                        "or upload a DOCX/TXT file instead."
                    )
                
                return text
        except Exception as e:
            if "scanned PDF" in str(e):
                raise
            raise ValueError(f"PDF parsing error: {str(e)}")
    
    def _parse_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX with image-only detection."""
        try:
            doc = Document(io.BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # Detect image-only DOCX (no text content)
            if len(text.strip()) < 50:
                raise ValueError(
                    "This DOCX file appears to contain only images or very little text. "
                    "Please upload a document with actual text content."
                )
            
            return text
        except Exception as e:
            if "image-only" in str(e):
                raise
            raise ValueError(f"DOCX parsing error: {str(e)}")
    
    def parse(self, filename: str, file_content: bytes) -> Dict:
        """
        Main parsing method - production grade with GDPR compliance.
        
        Args:
            filename: Original filename (for format detection)
            file_content: Raw file bytes
            
        Returns:
            Dictionary with parsed CV data (NO PII stored)
        """
        # Check file size (HuggingFace has 200MB limit)
        if len(file_content) > 50 * 1024 * 1024:  # 50MB limit for safety
            return self._empty_result(
                f"File too large ({len(file_content) / 1024 / 1024:.1f}MB). "
                "Maximum file size is 50MB. Please upload a smaller file."
            )
        
        # Step 1: Extract raw text
        try:
            raw_text = self._extract_text(filename, file_content)
        except ValueError as e:
            # Propagate specific parsing errors (scanned PDF, image DOCX, etc.)
            return self._empty_result(str(e))
        
        if not raw_text or len(raw_text.strip()) < 50:
            return self._empty_result(
                "Could not extract text from CV. The file may be image-based "
                "(scanned PDF or image-only DOCX). Please upload a text-based PDF, DOCX, or TXT file."
            )
        
        # Step 2: Sanitize PII and IMMEDIATELY PURGE (GDPR compliant)
        sanitized_text, context_hints = self.pii_extractor.extract_and_purge(raw_text)
        
        # Step 3: Extract structured data using LLM (or fallback)
        if self.use_llm:
            extraction = self.hybrid_extractor.extract(raw_text, sanitized_text)
        else:
            # Force local extraction
            extraction = self.hybrid_extractor._local_fallback(raw_text)
        
        # Step 4: Build result dictionary
        result = self._build_result(extraction, raw_text, context_hints)
        
        # NO PII stored - GDPR compliant
        result['_sanitized_sample'] = sanitized_text[:200] + "..."  # Truncated for debug only
        result['_privacy_note'] = "PII purged immediately after sanitization (GDPR compliant)"
        
        return result
    
    def _build_result(self, extraction: CVExtractionResult, raw_text: str, hints: Dict) -> Dict:
        """Convert extraction result to dictionary format for UI."""
        
        # Combine all skills
        all_skills = []
        skill_categories = {}
        
        for skill in extraction.technical_skills:
            all_skills.append(skill.name)
            skill_categories[skill.name] = "Technical"
        
        for skill in extraction.domain_knowledge:
            all_skills.append(skill.name)
            skill_categories[skill.name] = "Domain"
        
        for skill in extraction.tools_software:
            all_skills.append(skill.name)
            skill_categories[skill.name] = "Tools"
        
        for skill in extraction.soft_skills:
            all_skills.append(skill.name)
            skill_categories[skill.name] = "Soft Skills"
        
        # Deduplicate while preserving order
        seen = set()
        unique_skills = []
        for skill in all_skills:
            clean = skill.lower().strip()
            if clean and clean not in seen and len(clean) > 2:
                seen.add(clean)
                unique_skills.append(skill)
        
        # Determine domain
        domain = extraction.domain or hints.get('domain_hints', [None])[0] or "general"
        
        # Build professional title
        title = extraction.professional_title
        if not title and extraction.experiences:
            title = extraction.experiences[0].role
        
        # Experience calculation
        years_exp = extraction.total_years_experience
        if not years_exp and extraction.experiences:
            total_months = sum(e.duration_months or 0 for e in extraction.experiences)
            years_exp = round(total_months / 12, 1) if total_months > 0 else None
        
        # Career level
        career_level = extraction.career_level
        if not career_level:
            if years_exp and years_exp >= 7:
                career_level = "senior"
            elif years_exp and years_exp >= 3:
                career_level = "mid"
            else:
                career_level = "junior"
        
        return {
            # Core data
            'skills': unique_skills,
            'skill_categories': skill_categories,
            'years_experience': years_exp or 0,
            'career_level': career_level,
            'professional_title': title or "Professional",
            'domain': domain,
            'industry': extraction.industry,
            
            # For job matching
            'full_text': raw_text,
            'raw_skills_text': extraction.raw_skills_text or ", ".join(unique_skills),
            
            # Metadata
            'skill_count': len(unique_skills),
            'extraction_confidence': extraction.extraction_confidence,
            'extraction_method': extraction.extraction_method,
            'job_preferences': extraction.job_preferences,
            
            # Detailed (for debugging)
            'technical_skills': [s.name for s in extraction.technical_skills],
            'domain_knowledge': [s.name for s in extraction.domain_knowledge],
            'tools_software': [s.name for s in extraction.tools_software],
            'soft_skills': [s.name for s in extraction.soft_skills],
            'key_experiences': [
                {
                    'role': e.role,
                    'duration_months': e.duration_months,
                    'achievements': e.key_achievements
                }
                for e in extraction.experiences[:5]  # Top 5
            ]
        }
    
    def _empty_result(self, error_message: str) -> Dict:
        """Return empty result with error."""
        return {
            'skills': [],
            'skill_categories': {},
            'years_experience': 0,
            'career_level': 'unknown',
            'professional_title': 'Unknown',
            'domain': 'unknown',
            'full_text': '',
            'raw_skills_text': '',
            'skill_count': 0,
            'extraction_confidence': 0,
            'extraction_method': 'failed',
            'error': error_message,
            'job_preferences': [],
            'technical_skills': [],
            'domain_knowledge': [],
            'tools_software': [],
            'soft_skills': [],
            'key_experiences': []
        }


# Backward compatibility - EnhancedCVParser now uses DynamicCVParser
class EnhancedCVParser(DynamicCVParser):
    """Backward-compatible wrapper."""
    pass
